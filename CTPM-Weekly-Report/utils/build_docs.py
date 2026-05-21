"""
Build CTPM Sales Flash and Weekly Update .docx files.

Brand colors:
  Maroon:     #8B0000
  Steel blue: #1f77b4
  Header bg:  #D9D9D9
  Accent bg:  #F2F2F2
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
MAROON = RGBColor(0x8B, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x77, 0xB4)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
HEADER_BG = "D9D9D9"
ACCENT_BG = "F2F2F2"
MAROON_HEX = "8B0000"
BLUE_HEX = "1F77B4"

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"

SCHEDULER_NOTE = (
    "Our scheduler was out for a few days at the start of the month, resulting in a lighter "
    "than typical field schedule. As our scheduler gets caught up, field job volume is expected "
    "to climb back toward our normal pace."
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _fmt_dollar(v) -> str:
    try:
        return f"${float(v):,.0f}"
    except (TypeError, ValueError):
        return "$0"


def _fmt_num(v) -> str:
    try:
        return f"{int(round(float(v))):,}"
    except (TypeError, ValueError):
        return "0"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, border_color="000000", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_run(para, text, bold=False, italic=False, size=11, color=None, font=FONT_BODY):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = FONT_HEADING
    run.bold = True
    run.font.color.rgb = MAROON
    size_map = {1: 16, 2: 13, 3: 11}
    run.font.size = Pt(size_map.get(level, 11))
    return para


def _add_body_para(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    _para_run(para, text, bold=bold, size=size)
    return para


def _add_bullet(doc, text, size=11):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    return para


def _style_header_row(table, col_count):
    """Apply maroon background + white bold text to first row."""
    row = table.rows[0]
    for i, cell in enumerate(row.cells):
        _set_cell_bg(cell, MAROON_HEX)
        _set_cell_borders(cell)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_BODY


def _style_data_rows(table):
    """Alternate light grey / white on data rows; apply borders."""
    for i, row in enumerate(table.rows[1:], start=1):
        bg = ACCENT_BG if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)


def _table_cell(table, row_idx, col_idx, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell = table.cell(row_idx, col_idx)
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    _para_run(para, str(text), bold=bold, size=10)


def _new_doc() -> Document:
    doc = Document()
    # Set margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    # Remove default styles' extra spacing
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    return doc


def _add_ctpm_header(doc, title: str, subtitle: str = ""):
    """Add branded header block."""
    # Company name line
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run("CTPM")
    run.font.name = FONT_HEADING
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = MAROON

    # Divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(0)
    p_div.paragraph_format.space_after = Pt(4)
    run_div = p_div.add_run("─" * 72)
    run_div.font.color.rgb = MAROON
    run_div.font.size = Pt(9)

    # Document title
    para2 = doc.add_paragraph()
    para2.paragraph_format.space_before = Pt(2)
    para2.paragraph_format.space_after = Pt(2)
    r = para2.add_run(title)
    r.font.name = FONT_HEADING
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = DARK_GRAY

    if subtitle:
        para3 = doc.add_paragraph()
        para3.paragraph_format.space_before = Pt(0)
        para3.paragraph_format.space_after = Pt(8)
        rs = para3.add_run(subtitle)
        rs.font.name = FONT_BODY
        rs.font.size = Pt(11)
        rs.font.color.rgb = DARK_GRAY

    # Second divider
    p_div2 = doc.add_paragraph()
    p_div2.paragraph_format.space_before = Pt(0)
    p_div2.paragraph_format.space_after = Pt(8)
    run_div2 = p_div2.add_run("─" * 72)
    run_div2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run_div2.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Revenue table helper
# ---------------------------------------------------------------------------

def _add_revenue_table(doc, revenue, field_data):
    rows_data = [
        ("Field Jobs — Confirmed", _fmt_dollar(revenue["field_confirmed"]),
         f"{field_data['confirmed_count']} jobs" if field_data else "—"),
        ("Field Jobs — Pending", _fmt_dollar(revenue["field_pending"]),
         f"{field_data['pending_count']} jobs" if field_data else "—"),
        ("Shop Work Orders", _fmt_dollar(revenue["shop_total"]), "In-progress & quoted"),
        ("Estimated WIP Value", _fmt_dollar(revenue["wip_value"]), "Tiered historic lookup"),
        ("TOTAL PROJECTED", _fmt_dollar(revenue["total"]), ""),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    table.style = "Table Grid"

    headers = ["Revenue Component", "Estimated Value", "Notes"]
    for i, h in enumerate(headers):
        _table_cell(table, 0, i, h, bold=True)

    for r_idx, (label, val, note) in enumerate(rows_data, start=1):
        bold = label.startswith("TOTAL")
        _table_cell(table, r_idx, 0, label, bold=bold)
        _table_cell(table, r_idx, 1, val, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _table_cell(table, r_idx, 2, note)
        # Highlight total row
        if bold:
            for cell in table.rows[r_idx].cells:
                _set_cell_bg(cell, "E8D5D5")
                _set_cell_borders(cell)

    # Set column widths
    widths = [Inches(2.8), Inches(1.5), Inches(2.2)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    _style_header_row(table, 3)
    _style_data_rows(table)
    return table


# ---------------------------------------------------------------------------
# TAT table helper
# ---------------------------------------------------------------------------

def _add_tat_table(doc, tat):
    data = [
        ("Receipt → Calibration", f"{tat['receive_to_cal']} days"),
        ("Calibration → QC", f"{tat['cal_to_qc']} days"),
        ("QC → Shipment", f"{tat['qc_to_ship']} days"),
        ("Total (Receipt → Ship)", f"{tat['total']} days"),
    ]
    table = doc.add_table(rows=len(data) + 1, cols=2)
    table.style = "Table Grid"
    _table_cell(table, 0, 0, "TAT Stage", bold=True)
    _table_cell(table, 0, 1, "Average (Business Days)", bold=True)

    for i, (stage, val) in enumerate(data, start=1):
        bold = "Total" in stage
        _table_cell(table, i, 0, stage, bold=bold)
        _table_cell(table, i, 1, val, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        if bold:
            for cell in table.rows[i].cells:
                _set_cell_bg(cell, "E8D5D5")
                _set_cell_borders(cell)

    for row in table.rows:
        row.cells[0].width = Inches(2.8)
        row.cells[1].width = Inches(1.8)

    _style_header_row(table, 2)
    _style_data_rows(table)

    if tat.get("sample_size"):
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(2)
        note.paragraph_format.space_after = Pt(8)
        _para_run(note, f"n = {tat['sample_size']} instruments (last 90 days)", italic=True, size=9, color=DARK_GRAY)


# ---------------------------------------------------------------------------
# Sales Flash narrative paragraph
# ---------------------------------------------------------------------------

def _build_narrative(report_date, revenue, field_data, wip_data, cal_data, tat, override_narrative="") -> str:
    if override_narrative:
        return override_narrative

    month_name = report_date.strftime("%B")
    date_str = report_date.strftime("%B %d, %Y")

    total_field_count = 0
    confirmed_count = 0
    pending_count = 0
    field_total = revenue["field_confirmed"] + revenue["field_pending"]
    confirmed_total = revenue["field_confirmed"]
    pending_total = revenue["field_pending"]

    if field_data:
        total_field_count = field_data["confirmed_count"] + field_data["pending_count"]
        confirmed_count = field_data["confirmed_count"]
        pending_count = field_data["pending_count"]

    wip_items = wip_data["item_count"] if wip_data else 0
    wip_wos = wip_data["wo_count"] if wip_data else 0
    wip_value = revenue["wip_value"]
    cal_proj = revenue["cal_projection"]
    total_tat = tat["total"]
    total_rev = revenue["total"]

    scheduler_note = ""
    if confirmed_count < 35 and field_data is not None:
        scheduler_note = f" {SCHEDULER_NOTE}"

    narrative = (
        f"As of {date_str}, CTPM projects total {month_name} revenue of approximately "
        f"{_fmt_dollar(total_rev)}. "
        f"Currently, {total_field_count} field calibration job{'s are' if total_field_count != 1 else ' is'} "
        f"scheduled totaling {_fmt_dollar(field_total)}, with {confirmed_count} "
        f"job{'s' if confirmed_count != 1 else ''} ({_fmt_dollar(confirmed_total)}) converted to active "
        f"work orders and {pending_count} job{'s' if pending_count != 1 else ''} "
        f"({_fmt_dollar(pending_total)}) pending approval.{scheduler_note} "
        f"In-lab operations currently have {wip_items} instrument{'s' if wip_items != 1 else ''} "
        f"across {wip_wos} work order{'s' if wip_wos != 1 else ''} in the shop, with an estimated "
        f"value of {_fmt_dollar(wip_value)}. "
        f"Based on a six-month historical average, we project approximately {_fmt_num(cal_proj)} "
        f"in-lab calibration{'s' if cal_proj != 1 else ''} for the month. "
        f"Current shop turnaround time averages {total_tat} business days from receipt to shipment."
    )
    return narrative


# ---------------------------------------------------------------------------
# PUBLIC: Build Sales Flash
# ---------------------------------------------------------------------------

def build_sales_flash(path: str, report_date: datetime, revenue: dict,
                      field_data, wip_data, cal_data, manual_inputs: dict):
    doc = _new_doc()
    tat = manual_inputs["tat"]

    week_label = _week_label(report_date, manual_inputs)
    _add_ctpm_header(doc, "Sales Flash", subtitle=f"Week of {week_label}")

    # ── Revenue Breakdown ──
    _add_heading(doc, "Projected Revenue Breakdown", level=2)
    _add_revenue_table(doc, revenue, field_data)
    doc.add_paragraph()

    # Cal projection note (not a $ line in the total, just a metric)
    if cal_data and cal_data.get("monthly_average", 0):
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(2)
        note.paragraph_format.space_after = Pt(8)
        _para_run(
            note,
            f"Projected in-lab calibrations for the month: {_fmt_num(cal_data['monthly_average'])} "
            f"(6-month historical average, {cal_data.get('months_used', '?')} months).",
            italic=True, size=10,
        )

    # ── TAT ──
    _add_heading(doc, "Shop Turnaround Time (TAT)", level=2)
    _add_tat_table(doc, tat)
    doc.add_paragraph()

    # ── Key Metrics ──
    _add_heading(doc, "Key Metrics", level=2)
    table_m = doc.add_table(rows=5, cols=2)
    table_m.style = "Table Grid"
    metrics = [
        ("Field Jobs Scheduled (Confirmed)", str(field_data["confirmed_count"]) if field_data else "—"),
        ("Field Jobs Scheduled (Pending)", str(field_data["pending_count"]) if field_data else "—"),
        ("WIP Instruments in Lab", f"{wip_data['item_count']} across {wip_data['wo_count']} WOs" if wip_data else "—"),
        ("Estimated WIP Value", _fmt_dollar(revenue["wip_value"])),
        ("Projected Monthly In-Lab Calibrations", _fmt_num(revenue["cal_projection"]) if revenue["cal_projection"] else "—"),
    ]
    for i, (label, val) in enumerate(metrics):
        _table_cell(table_m, i, 0, label, bold=True)
        _table_cell(table_m, i, 1, val)
    for row in table_m.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(2.0)
    _style_data_rows(table_m)
    doc.add_paragraph()

    # ── Narrative Paragraph ──
    _add_heading(doc, "Summary", level=2)
    narrative = _build_narrative(
        report_date, revenue, field_data, wip_data, cal_data, tat,
        override_narrative=manual_inputs.get("sales_flash_narrative", ""),
    )
    body = doc.add_paragraph()
    body.paragraph_format.space_before = Pt(2)
    body.paragraph_format.space_after = Pt(4)
    _para_run(body, narrative, size=11)

    doc.save(path)


# ---------------------------------------------------------------------------
# PUBLIC: Build Weekly Update
# ---------------------------------------------------------------------------

def build_weekly_update(path: str, report_date: datetime, revenue: dict,
                        field_data, wip_data, cal_data, manual_inputs: dict):
    doc = _new_doc()
    tat = manual_inputs["tat"]

    week_start = manual_inputs.get("week_start", "")
    week_end = manual_inputs.get("week_end", "")
    if week_start and week_end:
        week_range = f"{week_start} – {week_end}"
    else:
        week_range = report_date.strftime("Week of %B %d, %Y")

    _add_ctpm_header(doc, "Weekly Operations Update", subtitle=week_range)

    # ── Operational Metrics ──
    _add_heading(doc, "Operational Metrics", level=2)

    ops_data = [
        ("WIP Instruments in Lab", f"{wip_data['item_count']}" if wip_data else "—"),
        ("Active Work Orders", f"{wip_data['wo_count']}" if wip_data else "—"),
        ("Estimated In-Lab WIP Value", _fmt_dollar(revenue["wip_value"])),
        ("Field Jobs Confirmed", str(field_data["confirmed_count"]) if field_data else "—"),
        ("Field Jobs Pending", str(field_data["pending_count"]) if field_data else "—"),
        ("Current Shop TAT (Receipt → Ship)", f"{tat['total']} business days"),
    ]
    table_ops = doc.add_table(rows=len(ops_data) + 1, cols=2)
    table_ops.style = "Table Grid"
    _table_cell(table_ops, 0, 0, "Metric", bold=True)
    _table_cell(table_ops, 0, 1, "Value", bold=True)
    for i, (label, val) in enumerate(ops_data, start=1):
        _table_cell(table_ops, i, 0, label)
        _table_cell(table_ops, i, 1, val)
    for row in table_ops.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(2.0)
    _style_header_row(table_ops, 2)
    _style_data_rows(table_ops)
    doc.add_paragraph()

    # ── TAT detail ──
    _add_heading(doc, "Turnaround Time Detail", level=2)
    _add_tat_table(doc, tat)
    doc.add_paragraph()

    # ── Key Achievements ──
    _add_heading(doc, "Key Achievements", level=2)
    achievements = manual_inputs.get("achievements", [])
    if achievements:
        for item in achievements:
            _add_bullet(doc, item)
    else:
        _add_body_para(doc, "(No achievements entered.)", size=10)
    doc.add_paragraph()

    # ── Challenges & Solutions ──
    _add_heading(doc, "Challenges & Solutions", level=2)
    challenges = manual_inputs.get("challenges", [])
    if challenges:
        for item in challenges:
            _add_bullet(doc, item)
    else:
        _add_body_para(doc, "(No challenges entered.)", size=10)
    doc.add_paragraph()

    # ── Innovations & Initiatives ──
    _add_heading(doc, "Innovations & Initiatives", level=2)
    innovations = manual_inputs.get("innovations", [])
    if innovations:
        for item in innovations:
            _add_bullet(doc, item)
    else:
        _add_body_para(doc, "(No innovations entered.)", size=10)
    doc.add_paragraph()

    # ── Customer Wins / Losses ──
    _add_heading(doc, "Customer Wins / Losses", level=2)
    wins = manual_inputs.get("customer_wins_losses", "").strip()
    wins_lines = [l.strip() for l in wins.splitlines() if l.strip()] if wins else []
    if wins_lines:
        for line in wins_lines:
            _add_bullet(doc, line)
    else:
        _add_body_para(doc, "(No entries.)")
    doc.add_paragraph()

    # ── Upcoming Week Focus ──
    _add_heading(doc, "Upcoming Week Focus", level=2)
    focus = manual_inputs.get("upcoming_focus", "").strip()
    focus_lines = [l.strip() for l in focus.splitlines() if l.strip()] if focus else []
    if focus_lines:
        for line in focus_lines:
            _add_bullet(doc, line)
    else:
        _add_body_para(doc, "(No entries.)")

    doc.save(path)


# ---------------------------------------------------------------------------
# Week label helper
# ---------------------------------------------------------------------------

def _week_label(report_date: datetime, manual_inputs: dict) -> str:
    start = manual_inputs.get("week_start", "").strip()
    end = manual_inputs.get("week_end", "").strip()
    if start and end:
        return f"{start} – {end}"
    if start:
        return start
    return report_date.strftime("%B %d, %Y")
